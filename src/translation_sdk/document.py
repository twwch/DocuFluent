import docx
from docx.document import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from dataclasses import dataclass, field
from typing import List, Optional, Union, Dict, Any
import logging
from copy import deepcopy
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

@dataclass
class TranslationSegment:
    id: str
    original_text: str
    translated_text: Optional[str] = None
    # We will store a reference to the actual object (Paragraph or Run or Cell) 
    # but since we can't easily pickle that, we might just keep it in memory during the process.
    # For the SDK, we'll assume the document object is kept alive.
    # For the SDK, we'll assume the document object is kept alive.
    obj_ref: Optional[object] = None 
    math_elements: Dict[str, Any] = field(default_factory=dict) 

class DocumentProcessor:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = docx.Document(file_path)
        self.segments: List[TranslationSegment] = []

    def _extract_text_and_math(self, para):
        text = ""
        math_elements = {}
        math_count = 0
        
        # Iterate over children of the paragraph element
        for child in para._element:
            if child.tag.endswith('r'): # Run
                # Extract text from run
                if child.text:
                    text += child.text
            elif child.tag.endswith('oMath') or child.tag.endswith('oMathPara'): # Inline Math
                key = f"{{{{MATH_{math_count}}}}}"
                math_elements[key] = child
                text += key
                math_count += 1
            else:
                # Other elements, try to extract text if possible
                if hasattr(child, 'itertext'):
                     for t in child.itertext():
                        text += t
        return text, math_elements

    def extract_segments(self) -> List[TranslationSegment]:
        """
        Iterates through the document and extracts text segments from paragraphs and tables.
        Skips empty segments.
        """
        self.segments = []
        visited_elements = set()

        def _process_container(container, prefix):
            # Extract from paragraphs
            for i, para in enumerate(container.paragraphs):
                if para._element in visited_elements:
                    continue
                visited_elements.add(para._element)
                
                text, math_elems = self._extract_text_and_math(para)
                text = text.strip()
                if text:
                    self.segments.append(TranslationSegment(
                        id=f"{prefix}p_{i}",
                        original_text=text,
                        obj_ref=para,
                        math_elements=math_elems
                    ))

            # Extract from tables
            for t_idx, table in enumerate(container.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        for p_idx, para in enumerate(cell.paragraphs):
                            if para._element in visited_elements:
                                continue
                            visited_elements.add(para._element)

                            text, math_elems = self._extract_text_and_math(para)
                            text = text.strip()
                            if text:
                                self.segments.append(TranslationSegment(
                                    id=f"{prefix}t_{t_idx}_r_{r_idx}_c_{c_idx}_p_{p_idx}",
                                    original_text=text,
                                    obj_ref=para,
                                    math_elements=math_elems
                                ))

        # Process Body
        _process_container(self.doc, "")

        # Process Headers and Footers
        for s_idx, section in enumerate(self.doc.sections):
            # Headers
            if section.header:
                _process_container(section.header, f"s_{s_idx}_h_")
            if section.first_page_header:
                _process_container(section.first_page_header, f"s_{s_idx}_h1_")
            if section.even_page_header:
                _process_container(section.even_page_header, f"s_{s_idx}_he_")
            
            # Footers
            if section.footer:
                _process_container(section.footer, f"s_{s_idx}_f_")
            if section.first_page_footer:
                _process_container(section.first_page_footer, f"s_{s_idx}_f1_")
            if section.even_page_footer:
                _process_container(section.even_page_footer, f"s_{s_idx}_fe_")
        
        logger.info(f"Extracted {len(self.segments)} segments from {self.file_path}")
        return self.segments

    def apply_translations(self, translations: Dict[str, str]):
        """
        Applies translations back to the document.
        translations: Dict mapping segment ID to translated text.
        """
        for seg in self.segments:
            if seg.id in translations:
                trans_text = translations[seg.id]
                
                # Skip if translation is identical to original (e.g. formulas, numbers)
                # This preserves the original paragraph structure (including formulas)
                if trans_text.strip() == seg.original_text.strip():
                    continue

                if seg.obj_ref and isinstance(seg.obj_ref, Paragraph):
                    para = seg.obj_ref
                    
                    # Get original text and run map
                    orig_text = ""
                    run_map = []
                    for run in para.runs:
                        for char in run.text:
                            orig_text += char
                            run_map.append(run)

                    # Clear existing runs
                    for run in para.runs:
                        run._element.getparent().remove(run._element)
                    
                    # Use SequenceMatcher to reconstruct paragraph with formatting
                    from difflib import SequenceMatcher
                    import re
                    
                    matcher = SequenceMatcher(None, orig_text, trans_text)
                    last_trans_pos = 0
                    
                    for match in matcher.get_matching_blocks():
                        a, b, size = match
                        
                        # 1. Handle unmatched translation text
                        if b > last_trans_pos:
                            unmatched_text = trans_text[last_trans_pos:b]
                            parts = re.split(r'(\{\{MATH_\d+\}\})', unmatched_text)
                            for part in parts:
                                if part in seg.math_elements:
                                    math_copy = deepcopy(seg.math_elements[part])
                                    para._element.append(math_copy)
                                else:
                                    if part:
                                        para.add_run(part)
                        
                        # 2. Handle matched text - Copy runs from original map
                        if size > 0:
                            current_run = None
                            current_text = ""
                            for k in range(a, a+size):
                                if k < len(run_map):
                                    run = run_map[k]
                                    if run != current_run:
                                        if current_run:
                                            new_run = para.add_run(current_text)
                                            self._copy_run_format(current_run, new_run)
                                        current_run = run
                                        current_text = ""
                                    current_text += orig_text[k]
                            if current_run:
                                new_run = para.add_run(current_text)
                                self._copy_run_format(current_run, new_run)
                                
                        last_trans_pos = b + size
                    
                    # Handle remaining unmatched translation
                    if last_trans_pos < len(trans_text):
                        unmatched_text = trans_text[last_trans_pos:]
                        parts = re.split(r'(\{\{MATH_\d+\}\})', unmatched_text)
                        for part in parts:
                            if part in seg.math_elements:
                                math_copy = deepcopy(seg.math_elements[part])
                                para._element.append(math_copy)
                            else:
                                if part:
                                    para.add_run(part)

    def _copy_run_format(self, source_run, target_run):
        """Copies font, size, color, and other formatting from source_run to target_run."""
        if not source_run or not target_run:
            return
            
        # Copy font attributes
        if source_run.font.name:
            target_run.font.name = source_run.font.name
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        if source_run.font.bold is not None:
            target_run.font.bold = source_run.font.bold
        if source_run.font.italic is not None:
            target_run.font.italic = source_run.font.italic
        if source_run.font.underline is not None:
            target_run.font.underline = source_run.font.underline
        if source_run.font.strike is not None:
            target_run.font.strike = source_run.font.strike
        if source_run.font.subscript is not None:
            target_run.font.subscript = source_run.font.subscript
        if source_run.font.superscript is not None:
            target_run.font.superscript = source_run.font.superscript
        if source_run.font.color.type:
             try:
                 target_run.font.color.rgb = source_run.font.color.rgb
             except:
                 pass # Theme colors might fail

    def save(self, output_path: str):
        self.doc.save(output_path)

    def save_bilingual(self, output_path: str, translations: Dict[str, str]):
        """
        Creates a bilingual document (Original / Translation).
        """
        doc_copy = docx.Document(self.file_path)
        from difflib import SequenceMatcher
        import re
        
        visited_elements = set()

        def _process_paragraph_bilingual(para, seg_id):
            if para._element in visited_elements:
                return
            visited_elements.add(para._element)

            if seg_id in translations:
                trans_data = translations[seg_id]
                
                # Handle both string (legacy) and dict (with score) formats
                if isinstance(trans_data, dict):
                    trans = trans_data.get("text", "")
                    score = trans_data.get("score", 10)
                else:
                    trans = trans_data
                    score = 10

                # Determine Color
                from docx.shared import RGBColor
                text_color = None
                if score < 5:
                    text_color = RGBColor(255, 0, 0) # Red
                elif score < 8.5:
                    text_color = RGBColor(255, 192, 0) # Orange/Yellow
                
                # Get original text and run map
                orig_text = ""
                run_map = []
                for run in para.runs:
                    for char in run.text:
                        orig_text += char
                        run_map.append(run)
                
                # Skip if identical (ignoring whitespace)
                if trans.strip() == orig_text.strip():
                    return
                    
                # Append new line
                para.add_run("\n")
                
                # Use SequenceMatcher to find common parts (formulas, numbers)
                matcher = SequenceMatcher(None, orig_text, trans)
                last_trans_pos = 0
                
                for match in matcher.get_matching_blocks():
                    a, b, size = match
                    # a: start in orig, b: start in trans, size: length
                    
                    # 1. Handle unmatched translation text before this match
                    if b > last_trans_pos:
                        unmatched_text = trans[last_trans_pos:b]
                        # Check for math placeholders in unmatched text
                        parts = re.split(r'(\{\{MATH_\d+\}\})', unmatched_text)
                        
                        # We need to find the segment to get math elements for placeholders
                        segment = next((s for s in self.segments if s.id == seg_id), None)
                        math_elems = segment.math_elements if segment else {}

                        for part in parts:
                            if part in math_elems:
                                math_copy = deepcopy(math_elems[part])
                                para._element.append(math_copy)
                            else:
                                if part:
                                    run = para.add_run(part)
                                    if text_color:
                                        run.font.color.rgb = text_color
                    
                    # 2. Handle matched text - Copy runs from original
                    if size > 0:
                        current_run = None
                        current_text = ""
                        
                        for k in range(a, a+size):
                            if k < len(run_map):
                                run = run_map[k]
                                if run != current_run:
                                    if current_run:
                                        new_run = para.add_run(current_text)
                                        self._copy_run_format(current_run, new_run)
                                        if text_color:
                                            new_run.font.color.rgb = text_color
                                    current_run = run
                                    current_text = ""
                                current_text += orig_text[k]
                        
                        # Flush last run
                        if current_run:
                            new_run = para.add_run(current_text)
                            self._copy_run_format(current_run, new_run)
                            if text_color:
                                new_run.font.color.rgb = text_color
                            
                    last_trans_pos = b + size
                
                # Handle remaining unmatched translation
                if last_trans_pos < len(trans):
                    unmatched_text = trans[last_trans_pos:]
                    # Check for math placeholders
                    parts = re.split(r'(\{\{MATH_\d+\}\})', unmatched_text)
                    segment = next((s for s in self.segments if s.id == seg_id), None)
                    math_elems = segment.math_elements if segment else {}

                    for part in parts:
                        if part in math_elems:
                            math_copy = deepcopy(math_elems[part])
                            para._element.append(math_copy)
                        else:
                            if part:
                                run = para.add_run(part)
                                if text_color:
                                    run.font.color.rgb = text_color

        def _process_container_bilingual(container, prefix):
            for i, para in enumerate(container.paragraphs):
                _process_paragraph_bilingual(para, f"{prefix}p_{i}")
            
            for t_idx, table in enumerate(container.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        for p_idx, para in enumerate(cell.paragraphs):
                            _process_paragraph_bilingual(para, f"{prefix}t_{t_idx}_r_{r_idx}_c_{c_idx}_p_{p_idx}")

        # Process Body
        _process_container_bilingual(doc_copy, "")

        # Process Headers and Footers
        for s_idx, section in enumerate(doc_copy.sections):
             # Headers
            if section.header:
                _process_container_bilingual(section.header, f"s_{s_idx}_h_")
            if section.first_page_header:
                _process_container_bilingual(section.first_page_header, f"s_{s_idx}_h1_")
            if section.even_page_header:
                _process_container_bilingual(section.even_page_header, f"s_{s_idx}_he_")
            
            # Footers
            if section.footer:
                _process_container_bilingual(section.footer, f"s_{s_idx}_f_")
            if section.first_page_footer:
                _process_container_bilingual(section.first_page_footer, f"s_{s_idx}_f1_")
            if section.even_page_footer:
                _process_container_bilingual(section.even_page_footer, f"s_{s_idx}_fe_")
        
        doc_copy.save(output_path)
