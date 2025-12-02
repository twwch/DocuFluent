import docx
import sys

def inspect_docx(file_path):
    doc = docx.Document(file_path)
    print(f"Inspecting {file_path}...")
    def check_para(para, label):
        if "28" in para.text or "MPa" in para.text:
            print(f"\n{label}: {para.text[:50]}...")
            for child in para._element:
                print(f"  Child Tag: {child.tag}")
                if child.tag.endswith('oMath') or child.tag.endswith('oMathPara'):
                    print("    -> FOUND MATH ELEMENT")
                if child.tag.endswith('r'):
                    print(f"    -> Run Text: {child.text}")

    for i, para in enumerate(doc.paragraphs):
        check_para(para, f"Paragraph {i}")

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    check_para(para, f"Table {t_idx} Row {r_idx} Cell {c_idx} Para {p_idx}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        inspect_docx(sys.argv[1])
    else:
        print("Please provide a file path.")
